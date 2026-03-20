from __future__ import annotations
"""Visual translation quality dashboard report."""

import math
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from docx import Document

from ..core.config import log
from ..docx.extractor import DocxExtractor
from ..translation.glossary import Glossary

_CJK_RE = re.compile(r"[\u4e00-\u9fff]")
_NUMBER_RE = re.compile(
    r"\d{1,3}(?:,\d{3})*(?:\.\d+)?|\d+(?:\.\d+)?%|\d{4}[-/]\d{1,2}[-/]\d{1,2}|[¥$€£]\d[\d,.]*"
)
_FORBIDDEN_EXPANSION_RE = re.compile(
    r"<think>|</think>|as an ai|translation:|explanation:|note:|总结|说明",
    re.IGNORECASE,
)


@dataclass
class QualityMetrics:
    name: str
    source_path: str
    target_path: str
    total_units: int
    aligned_units: int
    nonempty_units: int
    paragraph_match: bool
    table_match: bool
    cell_match: bool
    number_integrity: float
    cjk_residue_ratio: float
    repetition_consistency: float
    glossary_hit_ratio: Optional[float]
    scope_adherence: float
    professionalism: float
    completeness: float
    structure_fidelity: float
    terminology_consistency: float
    accuracy: float
    overall: float
    findings: List[str]


def _clamp(value: float, low: float = 0.0, high: float = 100.0) -> float:
    return max(low, min(high, value))


def _count_table_cells(doc: Document) -> int:
    return sum(len(row.cells) for table in doc.tables for row in table.rows)


def _extract_structure(path: str) -> Dict[str, int]:
    doc = Document(path)
    return {
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "table_cells": _count_table_cells(doc),
    }


def _extract_path_map(units) -> Dict[str, str]:
    path_map = {}
    for u in units:
        if u.source_text and u.source_text.strip():
            path_map[u.path] = u.source_text.strip()
    return path_map


def _extract_numbers(text: str) -> List[str]:
    return _NUMBER_RE.findall(text or "")


def _normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "")).strip().lower()


def _score_bar(score: float, width: int = 24, fill: str = "█", empty: str = "·") -> str:
    score = _clamp(score)
    filled = int(round(score / 100.0 * width))
    return fill * filled + empty * (width - filled)


def _resolve_targets(source_path: str, target_patterns: Optional[List[str]], output_dir: str) -> List[str]:
    if target_patterns:
        resolved = []
        for pattern in target_patterns:
            if any(ch in pattern for ch in ["*", "?", "["]):
                resolved.extend(str(p) for p in sorted(Path().glob(pattern)))
            else:
                resolved.append(pattern)
        return sorted(set(resolved))

    stem = Path(source_path).stem
    out_dir = Path(output_dir)
    auto = sorted(out_dir.glob(f"{stem}.en*.docx"))
    return [str(p) for p in auto]


def evaluate_pair(source_path: str, target_path: str, glossary: Optional[Glossary] = None) -> QualityMetrics:
    src_units = DocxExtractor(source_path).extract()
    tgt_units = DocxExtractor(target_path).extract()

    src_map = _extract_path_map(src_units)
    tgt_map = _extract_path_map(tgt_units)
    src_structure = _extract_structure(source_path)
    tgt_structure = _extract_structure(target_path)

    aligned = 0
    nonempty = 0
    number_checks = 0
    number_pass = 0
    cjk_units = 0
    scope_scores = []
    repetition_map = {}
    findings = []
    glossary_hit = 0
    glossary_total = 0
    forbidden_hits = 0

    for path, src_text in src_map.items():
        tgt_text = tgt_map.get(path, "")
        if path in tgt_map:
            aligned += 1
        if tgt_text.strip():
            nonempty += 1

        src_nums = _extract_numbers(src_text)
        if src_nums:
            number_checks += 1
            tgt_nums = _extract_numbers(tgt_text)
            ok = True
            for num in src_nums:
                n = num.replace(",", "")
                if not any(n in t.replace(",", "") for t in tgt_nums):
                    ok = False
                    break
            if ok:
                number_pass += 1

        if _CJK_RE.search(tgt_text):
            cjk_units += 1

        src_len = max(len(src_text), 1)
        tgt_len = len(tgt_text)
        ratio = tgt_len / float(src_len)
        if ratio < 0.25:
            scope_scores.append(35.0)
        elif ratio < 0.45:
            scope_scores.append(65.0)
        elif ratio <= 2.2:
            scope_scores.append(100.0)
        elif ratio <= 3.0:
            scope_scores.append(70.0)
        else:
            scope_scores.append(35.0)

        if _FORBIDDEN_EXPANSION_RE.search(tgt_text):
            forbidden_hits += 1

        key = _normalize_text(src_text)
        if key:
            repetition_map.setdefault(key, set()).add(_normalize_text(tgt_text))

        if glossary and glossary.entries:
            for entry in glossary.entries:
                if not entry.active:
                    continue
                if entry.source_term.lower() in src_text.lower():
                    glossary_total += 1
                    if entry.target_term.lower() in tgt_text.lower():
                        glossary_hit += 1

    total_units = len(src_map)
    structure_checks = [
        src_structure["paragraphs"] == tgt_structure["paragraphs"],
        src_structure["tables"] == tgt_structure["tables"],
        src_structure["table_cells"] == tgt_structure["table_cells"],
    ]
    paragraph_match, table_match, cell_match = structure_checks
    structure_fidelity = 100.0 * (sum(1 for ok in structure_checks if ok) / len(structure_checks))
    structure_fidelity = (structure_fidelity * 0.5) + (50.0 * (aligned / max(total_units, 1)))

    completeness = 100.0 * ((aligned / max(total_units, 1)) * 0.7 + (nonempty / max(total_units, 1)) * 0.3)

    number_integrity = 100.0 if number_checks == 0 else 100.0 * (number_pass / float(number_checks))
    cjk_ratio = cjk_units / float(max(aligned, 1))
    cjk_penalty = min(35.0, cjk_ratio * 100.0 * 0.45)

    repeat_total = 0
    repeat_stable = 0
    for src_text, variants in repetition_map.items():
        if src_text and len(variants) > 1:
            repeat_total += 1
        elif src_text and len(variants) == 1:
            repeat_total += 1
            repeat_stable += 1
    repetition_consistency = 100.0 if repeat_total == 0 else 100.0 * (repeat_stable / float(repeat_total))

    if glossary_total > 0:
        glossary_ratio = glossary_hit / float(glossary_total)
        glossary_hit_ratio = 100.0 * glossary_ratio
    else:
        glossary_hit_ratio = None

    scope_adherence = sum(scope_scores) / float(max(len(scope_scores), 1))
    scope_adherence = _clamp(scope_adherence - min(30.0, forbidden_hits * 4.0))

    professionalism = 100.0
    professionalism -= min(35.0, forbidden_hits * 6.0)
    professionalism -= cjk_penalty
    professionalism = _clamp(professionalism)

    terminology_consistency = repetition_consistency
    if glossary_hit_ratio is not None:
        terminology_consistency = (repetition_consistency * 0.55) + (glossary_hit_ratio * 0.45)

    accuracy = (
        number_integrity * 0.38
        + terminology_consistency * 0.32
        + scope_adherence * 0.12
        + (100.0 - cjk_penalty) * 0.18
    )
    accuracy = _clamp(accuracy)

    overall = (
        accuracy * 0.27
        + completeness * 0.18
        + structure_fidelity * 0.2
        + professionalism * 0.13
        + terminology_consistency * 0.12
        + scope_adherence * 0.1
    )
    overall = _clamp(overall)

    if not paragraph_match:
        findings.append("段落总数不一致，可能存在结构偏移。")
    if not table_match:
        findings.append("表格总数不一致，结构保真存在风险。")
    if not cell_match:
        findings.append("表格单元格总数不一致，可能出现内容丢失或错位。")
    if cjk_ratio > 0.05:
        findings.append("译文残留中文比例较高，存在漏译。")
    if forbidden_hits > 0:
        findings.append("检测到解释性或模型痕迹文本（如 <think>/note）。")
    if number_integrity < 95:
        findings.append("数字一致性偏低，建议重点复核数字与单位。")
    if glossary_hit_ratio is not None and glossary_hit_ratio < 90:
        findings.append("术语命中率不足，建议加强 glossary 约束。")

    return QualityMetrics(
        name=Path(target_path).name,
        source_path=source_path,
        target_path=target_path,
        total_units=total_units,
        aligned_units=aligned,
        nonempty_units=nonempty,
        paragraph_match=paragraph_match,
        table_match=table_match,
        cell_match=cell_match,
        number_integrity=_clamp(number_integrity),
        cjk_residue_ratio=cjk_ratio * 100.0,
        repetition_consistency=_clamp(repetition_consistency),
        glossary_hit_ratio=glossary_hit_ratio,
        scope_adherence=_clamp(scope_adherence),
        professionalism=_clamp(professionalism),
        completeness=_clamp(completeness),
        structure_fidelity=_clamp(structure_fidelity),
        terminology_consistency=_clamp(terminology_consistency),
        accuracy=_clamp(accuracy),
        overall=_clamp(overall),
        findings=findings,
    )


def render_dashboard(source_path: str, metrics_list: List[QualityMetrics]) -> str:
    lines = []
    lines.append("Translation Quality Intelligence Dashboard")
    lines.append("=" * 86)
    lines.append("Theme: Technical CN->EN Translation Quality Audit")
    lines.append("Source: {0}".format(source_path))
    lines.append("")

    ranked = sorted(metrics_list, key=lambda m: m.overall, reverse=True)
    lines.append("Model Ranking")
    lines.append("-" * 86)
    for idx, m in enumerate(ranked, 1):
        lines.append(
            "{0:>2}. {1:<50}  {2:6.2f}  {3}".format(
                idx, m.name[:50], m.overall, _score_bar(m.overall, width=16)
            )
        )
    lines.append("")

    for m in ranked:
        lines.append("┌─ {0}".format(m.name))
        lines.append("│  Overall:       {0:6.2f}   {1}".format(m.overall, _score_bar(m.overall)))
        lines.append("│  Accuracy:      {0:6.2f}   {1}".format(m.accuracy, _score_bar(m.accuracy)))
        lines.append("│  Completeness:  {0:6.2f}   {1}".format(m.completeness, _score_bar(m.completeness)))
        lines.append("│  Structure:     {0:6.2f}   {1}".format(m.structure_fidelity, _score_bar(m.structure_fidelity)))
        lines.append("│  Professional:  {0:6.2f}   {1}".format(m.professionalism, _score_bar(m.professionalism)))
        lines.append("│  Terminology:   {0:6.2f}   {1}".format(m.terminology_consistency, _score_bar(m.terminology_consistency)))
        lines.append("│  Scope Adhere:  {0:6.2f}   {1}".format(m.scope_adherence, _score_bar(m.scope_adherence)))
        lines.append("│")
        lines.append(
            "│  Units: aligned/nonempty/total = {0}/{1}/{2}".format(
                m.aligned_units, m.nonempty_units, m.total_units
            )
        )
        lines.append(
            "│  Structure checks: paragraph={0}, table={1}, cell={2}".format(
                "OK" if m.paragraph_match else "FAIL",
                "OK" if m.table_match else "FAIL",
                "OK" if m.cell_match else "FAIL",
            )
        )
        lines.append(
            "│  Number integrity={0:.2f}, CJK residue={1:.2f}%".format(
                m.number_integrity, m.cjk_residue_ratio
            )
        )
        if m.glossary_hit_ratio is not None:
            lines.append("│  Glossary hit ratio={0:.2f}%".format(m.glossary_hit_ratio))
        lines.append("│")
        if m.findings:
            lines.append("│  Findings:")
            for finding in m.findings[:5]:
                lines.append("│   - {0}".format(finding))
        else:
            lines.append("│  Findings: none")
        lines.append("└" + "─" * 84)
        lines.append("")

    return "\n".join(lines).rstrip() + "\n"


def generate_quality_dashboard(
    source_path: str,
    target_patterns: Optional[List[str]] = None,
    output_dir: str = "data/output",
    report_path: str = "",
    glossary_path: str = "",
) -> Tuple[str, List[QualityMetrics]]:
    source_p = Path(source_path)
    if not source_p.exists():
        raise FileNotFoundError("Source document not found: {0}".format(source_path))

    targets = _resolve_targets(source_path, target_patterns, output_dir)
    if not targets:
        raise ValueError("No target documents found. Provide --targets or ensure output docs exist.")

    glossary = None
    if glossary_path:
        glossary = Glossary()
        glossary.load_yaml(glossary_path)

    metrics_list = [evaluate_pair(source_path, target, glossary) for target in targets]
    dashboard = render_dashboard(source_path, metrics_list)

    if report_path:
        out = Path(report_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        with open(out, "w") as f:
            f.write("# Translation QA Dashboard\n\n")
            f.write("```text\n")
            f.write(dashboard)
            f.write("```\n")
        log.info("Quality dashboard written: {0}".format(out))
        return str(out), metrics_list

    return dashboard, metrics_list
