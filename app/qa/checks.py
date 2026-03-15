from __future__ import annotations
"""QA modules - automatic quality checks."""

import re
from typing import Optional

from ..core import TranslationUnit, QaIssue
from ..core.config import log


def extract_numbers(text: str) -> list[str]:
    """Extract all numbers, percentages, dates, currency amounts from text."""
    patterns = [
        r'\d{1,3}(?:,\d{3})*(?:\.\d+)?',  # 1,234.56
        r'\d+(?:\.\d+)?%',                   # 99.99%
        r'\d{4}[-/]\d{1,2}[-/]\d{1,2}',     # 2026-03-15
        r'[¥$€£]\d[\d,.]*',                  # $1,234.56
        r'\d+(?:\.\d+)?\s*(?:ms|s|GB|MB|KB|TPS|QPS)',  # units
        r'\d+(?:\.\d+)?\s*(?:个|条|项|次|分|秒|小时|天|周|月|年)',  # Chinese units
        r'SLA\s*\d+',                         # SLA references
    ]
    numbers = []
    for pat in patterns:
        numbers.extend(re.findall(pat, text))
    return numbers


class StructureChecker:
    """Checks structural alignment between source and translated units."""

    def check(
        self,
        source_units: list[TranslationUnit],
        source_stats: dict,
        target_stats: Optional[dict] = None,
    ) -> list[QaIssue]:
        issues = []
        target_stats = target_stats or {}

        # Paragraph count
        src_para = source_stats.get("paragraphs", 0)
        tgt_para = target_stats.get("paragraphs", 0)
        if tgt_para and src_para != tgt_para:
            issues.append(QaIssue(
                severity="error",
                category="structure",
                detail=f"Paragraph count mismatch: source={src_para}, target={tgt_para}",
            ))

        # Table count
        src_tables = source_stats.get("tables", 0)
        tgt_tables = target_stats.get("tables", 0)
        if tgt_tables and src_tables != tgt_tables:
            issues.append(QaIssue(
                severity="error",
                category="structure",
                detail=f"Table count mismatch: source={src_tables}, target={tgt_tables}",
            ))

        # Unit count
        if len(source_units) == 0:
            issues.append(QaIssue(
                severity="error",
                category="structure",
                detail="No translatable units found in source document",
            ))

        return issues


class NumberChecker:
    """Checks number consistency between source and translated text."""

    def check(self, unit: TranslationUnit) -> list[QaIssue]:
        if not unit.translated_text or unit.error:
            return []

        issues = []
        src_nums = extract_numbers(unit.source_text)
        tgt_nums = extract_numbers(unit.translated_text)

        # Check for missing numbers
        for num in src_nums:
            # Normalize for comparison
            normalized = num.replace(",", "").replace(" ", "")
            found = any(
                normalized.replace(",", "").replace(" ", "") in t.replace(",", "").replace(" ", "")
                for t in tgt_nums
            )
            if not found:
                issues.append(QaIssue(
                    severity="error",
                    category="number",
                    unit_id=unit.unit_id,
                    detail=f"Number may be missing or changed: '{num}'",
                    source_snippet=unit.source_text[:200],
                    target_snippet=unit.translated_text[:200],
                ))

        return issues


class TermChecker:
    """Checks terminology consistency for a single unit against glossary."""

    def check_unit(self, unit: TranslationUnit, glossary: dict) -> list[QaIssue]:
        """Check a single unit against glossary."""
        if not unit.translated_text or unit.error:
            return []

        issues = []
        for src_term, expected_tgt in glossary.items():
            # Check if source term is in the source text
            if src_term.lower() not in unit.source_text.lower():
                continue
            # Check if expected translation is in the translated text
            if expected_tgt.lower() not in unit.translated_text.lower():
                issues.append(QaIssue(
                    severity="warning",
                    category="term",
                    unit_id=unit.unit_id,
                    detail=f"Term '{src_term}' expected as '{expected_tgt}' but not found in translation",
                    source_snippet=unit.source_text[:200],
                    target_snippet=unit.translated_text[:200],
                ))
        return issues


class MissingTranslationChecker:
    """Detects potentially untranslated segments."""

    CJK_RE = re.compile(r'[\u4e00-\u9fff]')

    def check(self, unit: TranslationUnit) -> list[QaIssue]:
        if not unit.translated_text or unit.error:
            return []

        issues = []

        # Check if translation still contains significant Chinese
        cjk_chars = len(self.CJK_RE.findall(unit.translated_text))
        total_chars = len(unit.translated_text)
        if total_chars > 0 and cjk_chars / total_chars > 0.3:
            issues.append(QaIssue(
                severity="error",
                category="missing",
                unit_id=unit.unit_id,
                detail=f"Translation may be incomplete: {cjk_chars}/{total_chars} chars are still Chinese",
                source_snippet=unit.source_text[:200],
                target_snippet=unit.translated_text[:200],
            ))

        # Any CJK residue in non-CJK target text is a warning
        if cjk_chars > 0 and total_chars > 20:
            cjk_ratio = cjk_chars / total_chars
            if cjk_ratio <= 0.3:
                # Small amount of Chinese residue — still worth flagging
                issues.append(QaIssue(
                    severity="warning",
                    category="cjk_residue",
                    unit_id=unit.unit_id,
                    detail=f"Residual Chinese detected: {cjk_chars} CJK chars in translated text",
                    source_snippet=unit.source_text[:200],
                    target_snippet=unit.translated_text[:200],
                ))

        # Check for suspiciously short translations
        src_len = len(unit.source_text)
        tgt_len = len(unit.translated_text)
        if src_len > 50 and tgt_len < src_len * 0.2:
            issues.append(QaIssue(
                severity="warning",
                category="missing",
                unit_id=unit.unit_id,
                detail=f"Translation suspiciously short: source={src_len} chars, target={tgt_len} chars",
            ))

        return issues


class ChineseResidueChecker:
    """Post-write check for any Chinese characters remaining in the output DOCX."""

    CJK_RE = re.compile(r'[\u4e00-\u9fff]')

    def check(self, output_path: str) -> list[QaIssue]:
        """Scan output DOCX for paragraphs containing Chinese characters."""
        try:
            from docx import Document
        except ImportError:
            return []

        try:
            doc = Document(output_path)
        except Exception as e:
            return [QaIssue(
                severity="error",
                category="cjk_residue",
                detail=f"Cannot open output file for Chinese residue check: {e}",
            )]

        issues = []
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text or len(text) < 3:
                continue
            cjk = self.CJK_RE.findall(text)
            if cjk:
                issues.append(QaIssue(
                    severity="warning" if len(cjk) <= 3 else "error",
                    category="cjk_residue",
                    unit_id=f"out_p[{i}]",
                    detail=f"Output paragraph [{i}] contains {len(cjk)} Chinese chars: {''.join(cjk[:10])}",
                    target_snippet=text[:120],
                ))

        # Also check table cells
        for ti, table in enumerate(doc.tables):
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    if not text or len(text) < 3:
                        continue
                    cjk = self.CJK_RE.findall(text)
                    if cjk:
                        issues.append(QaIssue(
                            severity="warning" if len(cjk) <= 3 else "error",
                            category="cjk_residue",
                            unit_id=f"out_t[{ti}]/r[{ri}]/c[{ci}]",
                            detail=f"Output table[{ti}] cell[{ri},{ci}] contains {len(cjk)} Chinese chars",
                            target_snippet=text[:120],
                        ))

        return issues
