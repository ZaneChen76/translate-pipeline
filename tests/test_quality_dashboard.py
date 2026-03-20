import unittest
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document

from app.qa.dashboard import evaluate_pair, generate_quality_dashboard
from app.qa.dashboard_image import render_dashboard_image


class QualityDashboardTest(unittest.TestCase):
    def _build_doc(self, path: Path, body_text: str):
        doc = Document()
        doc.add_paragraph(body_text)
        table = doc.add_table(rows=1, cols=1)
        table.rows[0].cells[0].text = "参数 100%"
        doc.save(str(path))

    def test_evaluate_pair_basic_scores(self):
        with TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            src = tmp_path / "s.docx"
            tgt = tmp_path / "s.en.good.docx"
            self._build_doc(src, "系统在2026年支持100%可用性")
            self._build_doc(tgt, "The system supports 100% availability in 2026.")

            metrics = evaluate_pair(str(src), str(tgt))
            self.assertGreater(metrics.overall, 60.0)
            self.assertEqual(metrics.total_units, 2)

    def test_generate_quality_dashboard_auto_discovery(self):
        with TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            inbox = tmp_path / "inbox"
            output = tmp_path / "output"
            inbox.mkdir(parents=True, exist_ok=True)
            output.mkdir(parents=True, exist_ok=True)

            src = inbox / "demo.docx"
            self._build_doc(src, "测试文本")

            tgt1 = output / "demo.en.a.docx"
            tgt2 = output / "demo.en.b.docx"
            self._build_doc(tgt1, "Test text A")
            self._build_doc(tgt2, "Test text B")

            dashboard_text, metrics = generate_quality_dashboard(
                source_path=str(src),
                output_dir=str(output),
            )
            self.assertIn("Model Ranking", dashboard_text)
            self.assertEqual(len(metrics), 2)

    def test_render_dashboard_image(self):
        with TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            src = tmp_path / "demo.docx"
            tgt = tmp_path / "demo.en.a.docx"
            png = tmp_path / "dashboard.png"
            self._build_doc(src, "测试文本 100%")
            self._build_doc(tgt, "Test text 100%")

            metrics = [evaluate_pair(str(src), str(tgt))]
            out = render_dashboard_image(str(src), metrics, str(png))
            self.assertEqual(out, str(png))
            self.assertTrue(png.exists())


if __name__ == "__main__":
    unittest.main()
