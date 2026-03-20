import unittest

from docx import Document

from app.core import TranslationUnit
from app.docx.writer import DocxWriter


class WriterTest(unittest.TestCase):
    def test_replaces_full_table_cell_without_residue(self):
        from pathlib import Path
        from tempfile import TemporaryDirectory

        with TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source = tmp_path / "src.docx"
            output = tmp_path / "out.docx"

            doc = Document()
            table = doc.add_table(rows=1, cols=1)
            cell = table.rows[0].cells[0]
            cell.paragraphs[0].text = "第一段中文"
            cell.add_paragraph("第二段中文")
            doc.save(str(source))

            unit = TranslationUnit(
                unit_id="u0001",
                part="table",
                path="table[0]/cell[0,0]",
                source_text="第一段中文\n第二段中文",
                translated_text="English merged translation",
            )

            writer = DocxWriter(str(source), [unit])
            writer.write(str(output))

            out_doc = Document(str(output))
            out_text = out_doc.tables[0].rows[0].cells[0].text
            self.assertIn("English merged translation", out_text)
            self.assertNotIn("中文", out_text)


if __name__ == "__main__":
    unittest.main()
